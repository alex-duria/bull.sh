"""Export tools for PDF and DOCX generation."""

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from bullsh.config import get_config
from bullsh.tools.base import ToolResult, ToolStatus
from bullsh.logging import log


# Check for optional dependencies
REPORTLAB_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    pass

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    pass


def _get_exports_dir() -> Path:
    """Get the exports directory, creating if needed."""
    config = get_config()
    exports_dir = config.data_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    return exports_dir


def _parse_markdown(content: str) -> list[dict[str, Any]]:
    """Parse markdown content into structured elements."""
    elements = []
    lines = content.split('\n')
    current_para = []

    for line in lines:
        # Headers
        if line.startswith('# '):
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            elements.append({'type': 'h1', 'text': line[2:].strip()})
        elif line.startswith('## '):
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            elements.append({'type': 'h2', 'text': line[3:].strip()})
        elif line.startswith('### '):
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            elements.append({'type': 'h3', 'text': line[4:].strip()})
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            elements.append({'type': 'bullet', 'text': line.strip()[2:]})
        # Table rows
        elif '|' in line and not line.strip().startswith('|--'):
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                elements.append({'type': 'table_row', 'cells': cells})
        # Empty line
        elif not line.strip():
            if current_para:
                elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
        # Regular text
        else:
            # Clean markdown formatting
            clean_line = line.strip()
            clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)  # Bold
            clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)  # Italic
            clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)  # Code
            if clean_line:
                current_para.append(clean_line)

    if current_para:
        elements.append({'type': 'paragraph', 'text': ' '.join(current_para)})

    return elements


async def export_to_pdf(
    content: str,
    filename: str | None = None,
    title: str | None = None,
) -> ToolResult:
    """
    Export content to PDF format.

    Args:
        content: Markdown content to export
        filename: Optional filename (auto-generated if not provided)
        title: Optional document title

    Returns:
        ToolResult with path to generated PDF
    """
    if not REPORTLAB_AVAILABLE:
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_to_pdf",
            error_message="reportlab not installed. Run: pip install bullsh[export]",
        )

    try:
        log("tools", "export_to_pdf: Generating PDF")

        exports_dir = _get_exports_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = filename or f"research_{timestamp}.pdf"
        if not filename.endswith('.pdf'):
            filename += '.pdf'
        filepath = exports_dir / filename

        # Create PDF
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
        ))
        styles.add(ParagraphStyle(
            name='CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
        ))
        styles.add(ParagraphStyle(
            name='CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceBefore=6,
            spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            name='CustomBullet',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            bulletIndent=10,
            spaceBefore=3,
            spaceAfter=3,
        ))

        # Parse content
        elements_data = _parse_markdown(content)
        story = []

        # Add title if provided
        if title:
            story.append(Paragraph(title, styles['CustomTitle']))
            story.append(Spacer(1, 12))

        # Add generation info
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | bullsh Investment Research",
            styles['Normal']
        ))
        story.append(Spacer(1, 20))

        # Process elements
        table_rows = []
        for elem in elements_data:
            if elem['type'] == 'h1':
                if table_rows:
                    story.append(_create_pdf_table(table_rows))
                    table_rows = []
                story.append(Paragraph(elem['text'], styles['CustomTitle']))
            elif elem['type'] == 'h2':
                if table_rows:
                    story.append(_create_pdf_table(table_rows))
                    table_rows = []
                story.append(Paragraph(elem['text'], styles['CustomH2']))
            elif elem['type'] == 'h3':
                if table_rows:
                    story.append(_create_pdf_table(table_rows))
                    table_rows = []
                story.append(Paragraph(elem['text'], styles['CustomH3']))
            elif elem['type'] == 'bullet':
                if table_rows:
                    story.append(_create_pdf_table(table_rows))
                    table_rows = []
                story.append(Paragraph(f"• {elem['text']}", styles['CustomBullet']))
            elif elem['type'] == 'paragraph':
                if table_rows:
                    story.append(_create_pdf_table(table_rows))
                    table_rows = []
                story.append(Paragraph(elem['text'], styles['CustomBody']))
            elif elem['type'] == 'table_row':
                table_rows.append(elem['cells'])

        # Flush remaining table
        if table_rows:
            story.append(_create_pdf_table(table_rows))

        # Build PDF
        doc.build(story)

        log("tools", f"export_to_pdf: Saved to {filepath}")

        return ToolResult(
            data={
                "path": str(filepath),
                "filename": filename,
                "format": "pdf",
            },
            confidence=1.0,
            status=ToolStatus.SUCCESS,
            tool_name="export_to_pdf",
        )

    except Exception as e:
        log("tools", f"export_to_pdf: Error - {e}", level="error")
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_to_pdf",
            error_message=str(e),
        )


def _create_pdf_table(rows: list[list[str]]) -> Table:
    """Create a PDF table from rows."""
    if not rows:
        return Spacer(1, 0)

    # Determine column widths
    num_cols = max(len(row) for row in rows)
    col_width = 6.5 * inch / num_cols

    # Pad rows to same length
    padded_rows = []
    for row in rows:
        padded = row + [''] * (num_cols - len(row))
        padded_rows.append(padded)

    table = Table(padded_rows, colWidths=[col_width] * num_cols)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    return table


async def export_to_docx(
    content: str,
    filename: str | None = None,
    title: str | None = None,
) -> ToolResult:
    """
    Export content to DOCX format.

    Args:
        content: Markdown content to export
        filename: Optional filename (auto-generated if not provided)
        title: Optional document title

    Returns:
        ToolResult with path to generated DOCX
    """
    if not DOCX_AVAILABLE:
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_to_docx",
            error_message="python-docx not installed. Run: pip install bullsh[export]",
        )

    try:
        log("tools", "export_to_docx: Generating DOCX")

        exports_dir = _get_exports_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = filename or f"research_{timestamp}.docx"
        if not filename.endswith('.docx'):
            filename += '.docx'
        filepath = exports_dir / filename

        # Create document
        doc = Document()

        # Add title
        if title:
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation info
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | bullsh Investment Research"
        )
        info_run.font.size = Pt(9)
        info_run.font.italic = True
        doc.add_paragraph()

        # Parse and add content
        elements = _parse_markdown(content)
        table_rows = []

        for elem in elements:
            if elem['type'] == 'h1':
                if table_rows:
                    _add_docx_table(doc, table_rows)
                    table_rows = []
                doc.add_heading(elem['text'], level=1)
            elif elem['type'] == 'h2':
                if table_rows:
                    _add_docx_table(doc, table_rows)
                    table_rows = []
                doc.add_heading(elem['text'], level=2)
            elif elem['type'] == 'h3':
                if table_rows:
                    _add_docx_table(doc, table_rows)
                    table_rows = []
                doc.add_heading(elem['text'], level=3)
            elif elem['type'] == 'bullet':
                if table_rows:
                    _add_docx_table(doc, table_rows)
                    table_rows = []
                doc.add_paragraph(elem['text'], style='List Bullet')
            elif elem['type'] == 'paragraph':
                if table_rows:
                    _add_docx_table(doc, table_rows)
                    table_rows = []
                doc.add_paragraph(elem['text'])
            elif elem['type'] == 'table_row':
                table_rows.append(elem['cells'])

        # Flush remaining table
        if table_rows:
            _add_docx_table(doc, table_rows)

        # Save document
        doc.save(str(filepath))

        log("tools", f"export_to_docx: Saved to {filepath}")

        return ToolResult(
            data={
                "path": str(filepath),
                "filename": filename,
                "format": "docx",
            },
            confidence=1.0,
            status=ToolStatus.SUCCESS,
            tool_name="export_to_docx",
        )

    except Exception as e:
        log("tools", f"export_to_docx: Error - {e}", level="error")
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_to_docx",
            error_message=str(e),
        )


def _add_docx_table(doc: "Document", rows: list[list[str]]) -> None:
    """Add a table to a DOCX document."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            # Bold header row
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()  # Add space after table


async def export_to_markdown(
    content: str,
    filename: str | None = None,
    title: str | None = None,
) -> ToolResult:
    """
    Export content to Markdown format.

    Args:
        content: Content to export
        filename: Optional filename (auto-generated if not provided)
        title: Optional document title

    Returns:
        ToolResult with path to generated Markdown file
    """
    try:
        log("tools", "export_to_markdown: Generating Markdown")

        exports_dir = _get_exports_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = filename or f"research_{timestamp}.md"
        if not filename.endswith('.md'):
            filename += '.md'
        filepath = exports_dir / filename

        # Build content with optional title
        output = ""
        if title:
            output += f"# {title}\n\n"
        output += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | bullsh Investment Research*\n\n"
        output += "---\n\n"
        output += content

        # Write file
        filepath.write_text(output, encoding='utf-8')

        log("tools", f"export_to_markdown: Saved to {filepath}")

        return ToolResult(
            data={
                "path": str(filepath),
                "filename": filename,
                "format": "markdown",
            },
            confidence=1.0,
            status=ToolStatus.SUCCESS,
            tool_name="export_to_markdown",
        )

    except Exception as e:
        log("tools", f"export_to_markdown: Error - {e}", level="error")
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_to_markdown",
            error_message=str(e),
        )


async def export_content(
    content: str,
    filename: str | None = None,
    format: str = "md",
    title: str | None = None,
) -> ToolResult:
    """
    Export content to specified format.

    Args:
        content: Content to export
        filename: Optional filename
        format: Output format (md, pdf, docx)
        title: Optional document title

    Returns:
        ToolResult with path to generated file
    """
    format = format.lower().lstrip('.')

    if format in ('md', 'markdown'):
        return await export_to_markdown(content, filename, title)
    elif format == 'pdf':
        return await export_to_pdf(content, filename, title)
    elif format in ('docx', 'doc'):
        return await export_to_docx(content, filename, title)
    else:
        return ToolResult(
            data={},
            confidence=0.0,
            status=ToolStatus.FAILED,
            tool_name="export_content",
            error_message=f"Unsupported format: {format}. Use md, pdf, or docx.",
        )
